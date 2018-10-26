#! /usr/bin/env python
import docx
from docx import Document
import xml.etree.ElementTree as ET
import uuid
import datetime
import tempfile

#Gets list of all .docx files in working directory. This assumes that 
#all docx files in the directory belong to the same experimental session 
#and would therefore be stored in the same XML record
def get_docs():
    eln_entries = []
    for docx_files in glob.glob("*.docx"):
        eln_entries.append(docx_files)
    return eln_entries        
        
#Extracts text notes from ELN Page
def extractText(doc):
    notes = ET.SubElement(acquisitionActivity, "notes")
    entry = ET.SubElement(notes, "entry",{"xsi:type":"nx:TextEntry"})
    for row in doc.paragraphs:
        p = ET.SubElement(entry, "p")
        p.text=("%s"%(row.text))
    tree = ET.ElementTree()
    tree._setroot(notes)
    return notes

#Extracts tables from ELN page and takes each cell value
#and structures the resulting XML tree how Carelyn Campbell 
#structured tabular data in her own XML records
def extractTable(doc):
    notes = ET.SubElement(acquisitionActivity, "notes")
    for iter_tables in doc.tables:
        entry = ET.SubElement(notes, "entry",{"xsi:type":"nx:TableEntry"})
        table = ET.SubElement(entry, "table")
        row_count=0
        row_id_count=0
        col_count=0
        header = ET.SubElement(table, "header")
        rows = ET.SubElement(table, "rows")
    
        for row_iter in iter_tables.rows:
            if row_count==0:
                for columns in row_iter.cells:
                    column = ET.SubElement(header, "column", id="%s"%(col_count))
                    column.text = columns.text
                    col_count+=1
                row_count+=1
                col_count=0
            if row_count!=0:
                row=ET.SubElement(rows, "row", seqno="%s"%(row_id_count))
                for columns in row_iter.cells:
                    column = ET.SubElement(row, "cell", colid="%s"%(col_count))
                    column.text="%s"%(columns.text)
                    col_count+=1
                row_id_count+=1
            col_count=0
    tree = ET.ElementTree()
    tree._setroot(notes)
    return tree        

#This function would return the tree to an acquisition activity element
#in the final XML record that is to be generated 
#(ie: record that contains calendar, image MD, and ELN MD)
def getNotes():
    docx_in_dir = get_docs()
    notes = ET.SubElement(acquisitionActivity, "notes", source="ELN")
    for file in docx_in_dir:
        doc = Document(file)
        textEntry = ET.SubElement(notes, "entry")
        textEntry = extractText(doc)
        tableEntry = ET.SubElement(notes, "entry")
        tableEntry = extractTable(doc)  
    tree = ET.ElementTree()
    tree._setroot(notes)
    return tree
